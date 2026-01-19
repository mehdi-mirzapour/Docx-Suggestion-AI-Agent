import asyncio
import base64
import os
import uuid
import logging
from pathlib import Path
from typing import Any

import logging

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("server")

from docx import Document
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Resource, Tool, TextContent, Prompt, ResourceTemplate
from pydantic import AnyUrl
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Initialize MCP server
app = Server("docx-editor-server")

# Storage for uploaded documents
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# In-memory storage for document metadata and suggestions
documents = {}
suggestions_store = {}


def extract_document_metadata(doc_path: str) -> dict:
    """Extract metadata from a Word document."""
    doc = Document(doc_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    return {
        "word_count": sum(len(p.split()) for p in paragraphs),
        "paragraph_count": len(paragraphs),
        "preview": paragraphs[0][:200] if paragraphs else "",
    }




def generate_suggestions(doc_path: str, request: str) -> list[dict]:
    """Generate AI-powered suggestions using GPT-4o-mini with batched processing."""
    doc = Document(doc_path)
    suggestions = []
    
    # Initialize OpenAI client
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or api_key == "your_openai_api_key_here":
        # Fallback to rule-based if no API key
        return generate_suggestions_fallback(doc_path, request)
    
    client = OpenAI(api_key=api_key)
    
    # Collect non-empty paragraphs with their indices
    paragraphs_to_process = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        text = paragraph.text
        
        # Skip very short paragraphs (less than 10 words)
        if len(text.split()) < 10:
            continue
        
        paragraphs_to_process.append((idx, text))
    
    # Batch paragraphs to reduce API calls (process 5 paragraphs at a time)
    BATCH_SIZE = 5
    
    for batch_start in range(0, len(paragraphs_to_process), BATCH_SIZE):
        batch = paragraphs_to_process[batch_start:batch_start + BATCH_SIZE]
        
        # Create a combined prompt with all paragraphs in the batch
        batch_text = "\n\n---PARAGRAPH SEPARATOR---\n\n".join(
            f"[PARAGRAPH {i}]\n{text}" 
            for i, (idx, text) in enumerate(batch)
        )
        
        try:
            # Call GPT-4o-mini for batch suggestions
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": f"""You are a professional document editor. Analyze the given paragraphs and suggest improvements based on this request: "{request}"

For each paragraph, return your response in this exact JSON format:
{{
    "suggestions": [
        {{
            "paragraph_number": 0,
            "has_suggestion": true/false,
            "suggested_text": "improved version of the text",
            "reason": "brief explanation of the change"
        }},
        ...
    ]
}}

Only suggest changes if they meaningfully improve the text. If no changes are needed for a paragraph, set has_suggestion to false for that paragraph.
Process all paragraphs provided and return suggestions for each one."""
                    },
                    {
                        "role": "user",
                        "content": batch_text
                    }
                ],
                temperature=0.3,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            # Parse AI response
            import json
            ai_response = json.loads(response.choices[0].message.content)
            
            # Extract suggestions for each paragraph in the batch
            batch_suggestions = ai_response.get("suggestions", [])
            
            for suggestion_data in batch_suggestions:
                paragraph_num = suggestion_data.get("paragraph_number", 0)
                
                # Map back to original paragraph index
                if paragraph_num < len(batch):
                    original_idx, original_text = batch[paragraph_num]
                    
                    if suggestion_data.get("has_suggestion", False):
                        suggestions.append({
                            "id": str(uuid.uuid4()),
                            "paragraph_index": original_idx,
                            "original": original_text,
                            "suggested": suggestion_data["suggested_text"],
                            "reason": suggestion_data["reason"],
                        })
        
        except Exception as e:
            # Log error but continue processing other batches
            print(f"Error processing batch starting at {batch_start}: {e}")
            continue
    
    return suggestions


def generate_suggestions_fallback(doc_path: str, request: str) -> list[dict]:
    """Fallback rule-based suggestions if OpenAI API is not available."""
    doc = Document(doc_path)
    suggestions = []
    
    # Simple rule-based suggestions for demonstration
    for idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
            
        text = paragraph.text
        
        # Example: Make text more formal
        if "more formal" in request.lower():
            if "don't" in text.lower():
                suggestions.append({
                    "id": str(uuid.uuid4()),
                    "paragraph_index": idx,
                    "original": text,
                    "suggested": text.replace("don't", "do not").replace("Don't", "Do not"),
                    "reason": "Replace contractions with full forms for formality",
                })
        
        # Example: Make text more concise
        if "concise" in request.lower() or "shorter" in request.lower():
            if len(text.split()) > 30:
                suggestions.append({
                    "id": str(uuid.uuid4()),
                    "paragraph_index": idx,
                    "original": text,
                    "suggested": " ".join(text.split()[:20]) + "...",
                    "reason": "Shorten long paragraph for conciseness",
                })
    
    return suggestions


def apply_changes_to_document(doc_path: str, selected_suggestions: list[dict]) -> str:
    """Apply selected suggestions to the document."""
    doc = Document(doc_path)
    
    # Sort suggestions by paragraph index in reverse to avoid index shifting
    sorted_suggestions = sorted(
        selected_suggestions, 
        key=lambda x: x["paragraph_index"], 
        reverse=True
    )
    
    for suggestion in sorted_suggestions:
        idx = suggestion["paragraph_index"]
        if idx < len(doc.paragraphs):
            # Replace paragraph text
            doc.paragraphs[idx].text = suggestion["suggested"]
            
            # Add comment to indicate change (Track Changes simulation)
            # Note: python-docx doesn't support true Track Changes,
            # so we'll add a comment or highlight instead
    
    # Save modified document
    output_path = doc_path.replace(".docx", "_modified.docx")
    doc.save(output_path)
    
    return output_path


@app.list_resources()
async def list_resources() -> list[Resource]:
    """List available resources (the widget)."""
    # Read the built widget HTML
    widget_path = Path("../frontend/dist/index.html")
    
    if widget_path.exists():
        widget_html = widget_path.read_text()
    else:
        # Fallback to a simple HTML if build doesn't exist
        widget_html = """
        <!DOCTYPE html>
        <html>
        <head><title>Document Editor</title></head>
        <body>
            <div id="root">Widget not built yet. Run 'npm run build' in frontend/</div>
        </body>
        </html>
        """
    
    return [
        Resource(
            uri="ui://widget/document-editor.html",
            name="Document Editor Widget",
            mimeType="text/html+skybridge",
            text=widget_html,
        )
    ]

@app.read_resource()
async def read_resource(uri: AnyUrl) -> str:
    """Read resource content."""
    if str(uri) == "ui://widget/document-editor.html":
        # Read the widget HTML (Same logic as list_resources, kept simple)
        widget_path = Path("../frontend/dist/index.html")
        if widget_path.exists():
            return widget_path.read_text()
        return """<!DOCTYPE html><html><body>Widget not built.</body></html>"""
    
    raise ValueError(f"Resource not found: {uri}")


@app.list_tools()
async def list_tools() -> list[Tool]:
    """List available tools."""
    return [
        Tool(
            name="upload_document",
            description="""Upload a Word document (.docx) for editing and analysis.

📋 WORKFLOW:
This tool accepts a .docx file via publicly accessible URL or base64 content.
After upload, it returns a doc_id that you'll use with analyze_document.

📥 INPUT OPTIONS:

Option 1 (RECOMMENDED): Use 'file_url'
- Provide a publicly accessible URL to the .docx file
- Best for files of any size
- Example: https://file.io/abc123, https://tmpfiles.org/xyz789

Option 2: Use 'content' (base64)
- Only for small files (<50KB)
- Must be complete, valid base64 encoding
- Large files may be truncated by the client

⚠️ IMPORTANT: You must provide EITHER 'file_url' OR 'content', not both.

✅ RETURNS:
- doc_id: Use this with analyze_document tool
- Document metadata: word count, paragraph count, preview""",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Name of the file (must end in .docx)",
                    },
                    "content": {
                        "type": "string",
                        "description": "Base64 encoded document content (optional if file_url is provided). Only use for small files <50KB.",
                    },
                    "file_url": {
                        "type": "string",
                        "description": "Publicly accessible URL to download the .docx file (RECOMMENDED). Use this for files of any size.",
                    },
                },
                "required": ["filename"],
                "additionalProperties": False,
            },
            annotations={
                "destructiveHint": False,
                "openWorldHint": False,
                "readOnlyHint": True,
            },
            _meta={
                "openai/toolInvocation/invoking": "📤 Uploading document...",
                "openai/toolInvocation/invoked": "✅ Document uploaded successfully",
                "openai/widgetAccessible": False,
            }
        ),
        Tool(
            name="analyze_document",
            description="""Analyze an uploaded Word document and generate AI-powered suggestions.

📋 PREREQUISITES:
You must first call upload_document to get a doc_id.

🔍 WHAT IT DOES:
- Analyzes document paragraphs using GPT-4o-mini
- Generates contextual suggestions based on your request
- Returns suggestions with original text, improved text, and reasoning

💡 EXAMPLE REQUESTS:
- "Make it more formal and professional"
- "Fix grammar and spelling errors"
- "Make it more concise"
- "Improve clarity and readability"

✅ RETURNS:
Interactive widget showing all suggestions with:
- Original paragraph text
- Suggested improvements
- Explanation of changes
- Accept/reject options""",
            inputSchema={
                "type": "object",
                "properties": {
                    "doc_id": {
                        "type": "string",
                        "description": "Document ID returned by upload_document tool",
                    },
                    "request": {
                        "type": "string",
                        "description": "Your editing request (e.g., 'Make it more formal', 'Fix grammar', 'Improve clarity')",
                    },
                },
                "required": ["doc_id", "request"],
                "additionalProperties": False,
            },
            annotations={
                "destructiveHint": False,
                "openWorldHint": False,
                "readOnlyHint": True,
            },
            _meta={
                "openai/outputTemplate": "ui://widget/document-editor.html",
                "openai/toolInvocation/invoking": "🔍 Analyzing document...",
                "openai/toolInvocation/invoked": "✅ Analysis complete",
                "openai/widgetAccessible": True,
                "openai/widget/csp": "default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline';",
                "openai/widget/domain": "beata-discriminantal-sirena.ngrok-free.dev"
            }
        ),
        Tool(
            name="apply_changes",
            description="""Apply selected suggestions to create a modified Word document.

📋 PREREQUISITES:
You must have:
1. Called upload_document to get a doc_id
2. Called analyze_document to generate suggestions

✏️ WHAT IT DOES:
- Applies the selected suggestions to the document
- Creates a new modified .docx file
- Provides a download link for the updated document

📥 INPUT:
- doc_id: The document identifier
- suggestion_ids: Array of suggestion IDs to apply (from analyze_document)

✅ RETURNS:
- Download URL for the modified document
- Count of applied changes
- Interactive widget with download button""",
            inputSchema={
                "type": "object",
                "properties": {
                    "doc_id": {
                        "type": "string",
                        "description": "Document ID returned by upload_document tool",
                    },
                    "suggestion_ids": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of suggestion IDs to apply (from analyze_document results)",
                    },
                },
                "required": ["doc_id", "suggestion_ids"],
                "additionalProperties": False,
            },
            annotations={
                "destructiveHint": False,
                "openWorldHint": False,
                "readOnlyHint": True,
            },
            _meta={
                "openai/outputTemplate": "ui://widget/document-editor.html",
                "openai/toolInvocation/invoking": "✏️ Applying changes...",
                "openai/toolInvocation/invoked": "✅ Changes applied successfully",
                "openai/widgetAccessible": True,
                "openai/widget/csp": "default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline';",
                "openai/widget/domain": "beata-discriminantal-sirena.ngrok-free.dev"
            }
        ),
    ]

@app.list_prompts()
async def list_prompts() -> list[Prompt]:
    """List available prompts."""
    return []

@app.list_resource_templates()
async def list_resource_templates() -> list[ResourceTemplate]:
    """List available resource templates."""
    return []


@app.call_tool()
async def call_tool(name: str, arguments: Any) -> list[TextContent]:
    """Handle tool calls."""
    
    if name == "upload_document":
        # Handle Upload Logic
        filename = arguments["filename"]
        
        # Get file content from either base64 or URL
        content = None
        
        if "file_url" in arguments and arguments["file_url"]:
            # Download from URL
            try:
                import httpx
                file_url = arguments["file_url"]
                logger.info(f"Downloading file from URL: {file_url}")
                
                response = httpx.get(file_url, follow_redirects=True, timeout=30.0)
                response.raise_for_status()
                content = response.content
                logger.info(f"Downloaded {len(content)} bytes from URL")
            except Exception as e:
                return [TextContent(type="text", text=f"Error downloading file from URL: {str(e)}")]
        
        elif "content" in arguments and arguments["content"]:
            # Decode from base64
            try:
                content_str = arguments["content"]
                
                # Detect truncation - suspiciously small base64 content
                if len(content_str) < 1000:
                    logger.warning(f"Received small base64 content ({len(content_str)} chars), likely truncated")
                    return [TextContent(
                        type="text",
                        text=f"⚠️ Warning: Received only {len(content_str)} characters of base64 data. "
                             f"This appears to be truncated.\n\n"
                             f"Please use 'file_url' parameter instead with a publicly accessible URL.\n\n"
                             f"Steps:\n"
                             f"1. Upload your file to a temporary hosting service (e.g., file.io, tmpfiles.org)\n"
                             f"2. Get the public download URL\n"
                             f"3. Call upload_document with file_url parameter"
                    )]
                
                if len(content_str) % 4:
                    content_str += '=' * (4 - len(content_str) % 4)
                content = base64.b64decode(content_str)
                logger.info(f"Decoded {len(content)} bytes from base64")
            except Exception as e:
                return [TextContent(type="text", text=f"Error decoding file content: {str(e)}")]
        else:
            # Neither content nor file_url provided
            return [TextContent(
                type="text",
                text="""❌ Error: Missing file content

You called upload_document with only 'filename', but you need to provide the actual file.

🔧 SOLUTION:

STEP 1: Upload the file to a temporary hosting service
You can use one of these services:
• file.io - Simple, no account needed
• tmpfiles.org - Alternative option
• Any other publicly accessible URL

STEP 2: Get the public download URL

STEP 3: Call upload_document again with both parameters:
{
  "filename": "Azure.docx",
  "file_url": "https://file.io/YOUR_FILE_ID"
}

📝 Example using file.io:
If you have access to curl or similar tools:
  curl -F "file=@Azure.docx" https://file.io
  
This returns: {"success":true,"link":"https://file.io/abc123"}
Then use "https://file.io/abc123" as the file_url parameter.

⚠️ Note: The file must be publicly accessible for download."""
            )]

        doc_id = str(uuid.uuid4())
        doc_path = UPLOAD_DIR / f"{doc_id}.docx"
        
        with open(doc_path, "wb") as f:
            f.write(content)
            
        # Verify ZIP/DOCX validity
        import zipfile
        if not zipfile.is_zipfile(doc_path):
             if doc_path.exists():
                doc_path.unlink()
             header_hex = content[:4].hex().upper()
             return [TextContent(type="text", text=f"Error: The uploaded file is not a valid DOCX/ZIP package. Header: {header_hex}, Size: {len(content)} bytes.")]

        # Extract metadata
        try:
            metadata = extract_document_metadata(str(doc_path))
        except Exception as e:
            if doc_path.exists():
                doc_path.unlink()
            return [TextContent(type="text", text=f"Error processing document structure: {str(e)}")]
        
        # Store document info
        documents[doc_id] = {
            "filename": filename,
            "path": str(doc_path),
            "metadata": metadata,
        }
        
        return [
            TextContent(
                type="text",
                text=f"Uploaded '{filename}' successfully.\n\nDocument ID: {doc_id}\nWord count: {metadata['word_count']}\nParagraphs: {metadata['paragraph_count']}"
            )
        ]
    
    elif name == "analyze_document":
        # Handle Analysis Logic
        doc_id = arguments["doc_id"]
        request = arguments["request"]
        
        if doc_id not in documents:
            return [TextContent(type="text", text="Document not found. Please upload the document first using upload_document.")]
        
        doc_path = documents[doc_id]["path"]
        filename = documents[doc_id]["filename"]
        
        # Generate suggestions
        suggestions = generate_suggestions(doc_path, request)
        suggestions_store[doc_id] = suggestions
        
        return [
            TextContent(
                type="text",
                text=f"Found {len(suggestions)} suggestions for: '{request}'",
                annotations={
                    "structuredContent": {
                        "doc_id": doc_id,
                        "filename": filename,
                        "suggestions": suggestions
                    }
                },
            )
        ]
    
    elif name == "apply_changes":
        doc_id = arguments["doc_id"]
        suggestion_ids = arguments["suggestion_ids"]
        
        if doc_id not in documents or doc_id not in suggestions_store:
            return [TextContent(type="text", text="Document or suggestions not found")]
        
        # Get selected suggestions
        all_suggestions = suggestions_store[doc_id]
        selected = [s for s in all_suggestions if s["id"] in suggestion_ids]
        
        # Apply changes
        doc_path = documents[doc_id]["path"]
        modified_path = apply_changes_to_document(doc_path, selected)
        
        # Create a user-friendly filename based on original filename
        original_filename = documents[doc_id]["filename"]
        # Remove .docx extension if present, add _modified, then add .docx
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f"{base_name}_modified.docx"

        # Store modified document path and download filename
        documents[doc_id]["modified_path"] = modified_path
        documents[doc_id]["download_filename"] = download_filename
        
        # Use valid public URL for download
        base_url = "https://beata-discriminantal-sirena.ngrok-free.dev"
        
        return [
            TextContent(
                type="text",
                text=f"Applied {len(selected)} changes to document",
                annotations={
                    "structuredContent": {
                        # Use public URL for the widget
                        "download_url": f"{base_url}/api/download/{doc_id}",
                        "applied_count": len(selected),
                    }
                },
            )
        ]
    
    return [TextContent(type="text", text=f"Unknown tool: {name}")]


from starlette.applications import Starlette
from starlette.middleware import Middleware
from starlette.middleware.cors import CORSMiddleware
from starlette.routing import Route, Mount
from mcp.server.sse import SseServerTransport
import uvicorn

# ... existing code ...

async def handle_sse(request):
    # Determine the absolute or relative path for the messages endpoint.
    # Since we are mounting the messages handler at /sse/messages, we should tell the client to post there.
    # SseServerTransport takes the endpoint URL that the client should POST to.
    transport = SseServerTransport("/sse/messages")
    async with app.run(transport.read_stream, transport.write_stream, app.create_initialization_options()) as streams:
        await streams
    return await transport.handle_sse(request)

async def handle_messages(request):
    # Since we are using Starlette, we need to bridge the request to the MCP transport.
    # However, SseServerTransport.handle_post_message expects scope/receive/send (ASGI) or similar.
    # The standard way with this SDK integration is tricky because `transport` is local to `handle_sse`.
    # BUT, actually, for the standalone SseServerTransport, we should instantiate it once globally or per-session?
    # No, MCP SDK design for Starlette/FastAPI usually implies we keep the transport alive.
    # Let's fix this by using the ASGI app pattern if possible or a global transport map.
    # FOR NOW: To keep it simple and working as per the previous working state (implied),
    # we will rely on the fact that handle_sse creates a session.
    # WAIT: Standard MCP python SDK usage for Starlette:
    
    # We need a shared transport handling mechanism.
    # Let's revert to a simpler "Global SSE" pattern for this demo to work reliably.
    
    pass

# Let's use the standard MCP SSE implementation pattern for Python
from starlette.responses import Response, JSONResponse

# SSE transport instance
# SSE transport instance
# We use the explicit path that matches our Mount structure to ensure the client
# receives the correct URL for POST requests (e.g. https://domain.com/sse/messages)
sse_transport = SseServerTransport("/sse/messages")

async def handle_mcp_sse(scope, receive, send):
    """
    Combined ASGI handler for MCP SSE (Connection & Messages).
    Mounted at /sse.
    """
import logging

# Configure Logging to File
logging.basicConfig(
    level=logging.DEBUG,
    filename="server.log",
    filemode="a",
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    force=True
)
logger = logging.getLogger("server")

# ... (rest of imports)

async def log_receive_wrapper(receive):
    """Wrapper to log incoming ASGI messages (request body)."""
    message = await receive()
    if message["type"] == "http.request":
        body = message.get("body", b"")
        if body:
            logger.debug(f"📥 POST Body: {body.decode('utf-8', errors='replace')}")
    return message

async def handle_mcp_sse(scope, receive, send):
    """
    Combined ASGI handler for MCP SSE (Connection & Messages).
    Mounted at /sse.
    """
    path = scope["path"]
    method = scope["method"]
    
    logger.debug(f"MCP Handler hit. Method: {method}, Path: {path}")

    if method == "POST":
         # Logic for message handling (POST /sse/messages)
         # We accept matching "/messages" or just loose matching if it's the only POST
         logger.debug("Handling POST message")
         try:
             # Wrap receive to capture body
             async def wrapped_receive():
                 message = await receive()
                 if message["type"] == "http.request":
                     body = message.get("body", b"")
                     if body:
                         logger.debug(f"📥 REQUEST BODY: {body.decode('utf-8', errors='replace')}")
                 return message
             
             # Actually, simpler: just read body if I can put it back? No.
             # Use the wrapper defined above? No, stateful closure needed for subsequent calls?
             # Standard ASGI receive is one-time or stream.
             # Let's simple-log inside a local wrapper.
             async def logging_receive():
                 msg = await receive()
                 if msg['type'] == 'http.request':
                     body = msg.get('body', b'')
                     logger.debug(f"📥 RAW RECEIVED: {body.decode('utf-8', errors='replace')}")
                 return msg

             await sse_transport.handle_post_message(scope, logging_receive, send)
         except Exception as e:
            logger.error(f"ERROR in handle_post_message: {e}")
            raise
    else:
         # Logic for connection (GET /sse)
         logger.debug("Handling SSE connection")
         async with sse_transport.connect_sse(scope, receive, send) as streams:
             await app.run(streams[0], streams[1], app.create_initialization_options())

async def handle_root(request):
    return JSONResponse({"status": "healthy", "message": "MCP Server is running"})

# REST API endpoints for standalone testing
from starlette.requests import Request
from starlette.responses import FileResponse

async def handle_upload(request: Request):
    """REST endpoint to upload a document."""
    form = await request.form()
    file = form.get("file")
    
    if not file:
        return JSONResponse({"error": "No file provided"}, status_code=400)
    
    # Read file content
    content = await file.read()
    filename = file.filename
    
    # Create doc_id and save
    doc_id = str(uuid.uuid4())
    doc_path = UPLOAD_DIR / f"{doc_id}.docx"
    
    with open(doc_path, "wb") as f:
        f.write(content)
    
    # Extract metadata
    metadata = extract_document_metadata(str(doc_path))
    
    # Store document info
    documents[doc_id] = {
        "filename": filename,
        "path": str(doc_path),
        "metadata": metadata,
    }
    
    return JSONResponse({
        "doc_id": doc_id,
        "filename": filename,
        "metadata": metadata
    })

async def handle_analyze(request: Request):
    """REST endpoint to analyze document and get suggestions."""
    data = await request.json()
    doc_id = data.get("doc_id")
    edit_request = data.get("request")
    
    if not doc_id or not edit_request:
        return JSONResponse({"error": "Missing doc_id or request"}, status_code=400)
    
    if doc_id not in documents:
        return JSONResponse({"error": "Document not found"}, status_code=404)
    
    doc_path = documents[doc_id]["path"]
    suggestions = generate_suggestions(doc_path, edit_request)
    
    # Store suggestions
    suggestions_store[doc_id] = suggestions
    
    return JSONResponse({
        "doc_id": doc_id,
        "suggestions": suggestions,
        "count": len(suggestions)
    })

async def handle_apply(request: Request):
    """REST endpoint to apply selected suggestions."""
    data = await request.json()
    doc_id = data.get("doc_id")
    suggestion_ids = data.get("suggestion_ids", [])
    
    if not doc_id:
        return JSONResponse({"error": "Missing doc_id"}, status_code=400)
    
    if doc_id not in documents or doc_id not in suggestions_store:
        return JSONResponse({"error": "Document or suggestions not found"}, status_code=404)
    
    # Get selected suggestions
    all_suggestions = suggestions_store[doc_id]
    selected = [s for s in all_suggestions if s["id"] in suggestion_ids]
    
    if not selected:
        return JSONResponse({"error": "No valid suggestions selected"}, status_code=400)
    
    # Apply changes
    doc_path = documents[doc_id]["path"]
    modified_path = apply_changes_to_document(doc_path, selected)
    
    # Create a user-friendly filename based on original filename
    original_filename = documents[doc_id]["filename"]
    # Remove .docx extension if present, add _modified, then add .docx
    base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
    download_filename = f"{base_name}_modified.docx"
    
    # Store modified document path and download filename
    documents[doc_id]["modified_path"] = modified_path
    documents[doc_id]["download_filename"] = download_filename
    
    return JSONResponse({
        "success": True,
        "applied_count": len(selected),
        "download_url": f"/api/download/{doc_id}"
    })

async def handle_download(request: Request):
    """REST endpoint to download modified document."""
    doc_id = request.path_params.get("doc_id")
    
    if doc_id not in documents:
        return JSONResponse({"error": "Document not found"}, status_code=404)
    
    modified_path = documents[doc_id].get("modified_path")
    if not modified_path or not Path(modified_path).exists():
        return JSONResponse({"error": "Modified document not found"}, status_code=404)
    
    # Use the stored download filename with proper extension
    download_filename = documents[doc_id].get("download_filename", "modified_document.docx")
    
    return FileResponse(
        path=str(modified_path),
        filename=download_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Define routes
routes = [
    Route("/", handle_root),
    Route("/api/upload", handle_upload, methods=["POST"]),
    Route("/api/analyze", handle_analyze, methods=["POST"]),
    Route("/api/apply", handle_apply, methods=["POST"]),
    Route("/api/download/{doc_id}", handle_download, methods=["GET"]),
    Mount("/sse", app=handle_mcp_sse),
]

# Configure CORS
middleware = [
    Middleware(
        CORSMiddleware,
        allow_origins=["*"],  # Allow all origins in development
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
        expose_headers=["*"],
    )
]

starlette_app = Starlette(debug=True, routes=routes, middleware=middleware)

if __name__ == "__main__":
    uvicorn.run(starlette_app, host="0.0.0.0", port=8787)

